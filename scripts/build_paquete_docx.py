#!/usr/bin/env python3
"""Assemble the consolidated "PAQUETE" .docx for one Constelaciones Familiares
publish-ready piece: image(s) + copy/caption + a new primer-comentario CTA +
hashtags (if the format uses them) + a fixed checklist, all in one file.

See PACKAGING_STANDARD in scripts/references/constelaciones_brand_voice.md
for the full rule this implements. Called by post-constelaciones,
carrusel-constelaciones, imagen-post-constelaciones, and
post-viral-constelaciones as the last step of their flow -- never by
historias-constelaciones or the reel skills (see that section for why).

This script never decides copy, CTA wording, or which book to mention --
those decisions are made by whoever is drafting the piece (the calling
skill), same division of labor as generate_post_image.py never writing the
caption it illustrates.

Usage:
  python scripts/build_paquete_docx.py "<hook o tema de la pieza>" \\
    --copy-docx "<ruta al .docx con el copy/caption ya aprobado>" \\
    --image "<ruta a imagen 1>" [--image "<ruta a imagen 2>" ...] \\
    [--primer-comentario "<texto corto, 2-3 líneas, CTA distinto al de la copy>"] \\
    --tipo-pieza <imagen-texto|carrusel|carrusel-historia|texto-reflexivo> \\
    --micronicho "<slug corto del micronicho, ej. dolor-heredado>" \\
    [--fecha "<YYYY-MM-DD, default hoy>"] \\
    [--out-dir "<override puntual, reemplaza --tipo-pieza/--micronicho/--fecha>"]

--image se puede repetir: 0 veces para post-viral-constelaciones (nunca
lleva imagen), 1 vez para post-constelaciones/imagen-post-constelaciones,
una vez por slide en orden de publicación para carrusel-constelaciones.

--primer-comentario es opcional: si se omite, esa sección del paquete no se
genera. Pensado para piezas sin CTA/libro por diseño (ej. un paquete
retroactivo de historias-constelaciones, que no usa PACKAGING_STANDARD de
forma automática pero puede empaquetarse a mano con este script) -- para
las 4 skills que sí llaman a este script como parte de su flujo normal,
--primer-comentario sigue siendo obligatorio en la práctica porque
PACKAGING_STANDARD lo exige.

Hashtags nunca se piden por separado -- se extraen automáticamente del
último párrafo del copy si empieza con "#" (la convención ya fija de dónde
van los hashtags en post-constelaciones/carrusel-constelaciones). Si el
copy no termina en un párrafo de hashtags (post-viral-constelaciones, que
nunca los lleva), esa sección del paquete se omite entera.

**Carpeta de salida.** Por defecto (sin --out-dir), el paquete se guarda en
`Desktop/Constelaciones - Publicaciones/<fecha> <micronicho>/<subcarpeta fija
de --tipo-pieza>/` -- ver PIECE_TYPE_SUBFOLDERS abajo para el mapeo exacto.
Una sola carpeta de día+micronicho junta TODAS las piezas de ese micronicho
(carrusel, post estático, virales), organizadas por tipo, sin importar cuál
de las 4 skills la generó. --tipo-pieza y --micronicho son obligatorios
salvo que se pase --out-dir explícitamente. Cuando un mismo micronicho tiene
varias piezas el mismo día (el caso normal -- un carrusel + un post + 2
virales), cada llamada debe pasar el MISMO --micronicho para que todas
caigan en la misma carpeta de día en vez de crear una por pieza; ver
PACKAGING_STANDARD para el criterio de cuándo reusar el slug de una pieza
hermana vs. derivar uno nuevo.
"""
from __future__ import annotations

import argparse
import datetime
import os
import sys
from pathlib import Path
from typing import Optional

import docx
from docx.shared import Inches

BASE_PUBLICACIONES_DIR = Path(os.path.expanduser("~")) / "Desktop" / "Constelaciones - Publicaciones"

# Fixed relative subfolder per piece type, inside <fecha> <micronicho>/ --
# same names/numbering PACKAGING_STANDARD documents, kept as constants here
# (not free text per call) so every skill produces byte-identical folder
# names instead of each one typing its own slightly different spelling.
# "carrusel-historia" has no calling skill yet (historias-constelaciones is
# not part of PACKAGING_STANDARD) -- it exists so a piece can still be
# packaged by hand into the right place, per PACKAGING_STANDARD's note on
# Historias.
PIECE_TYPE_SUBFOLDERS = {
    "imagen-texto": Path("Paquete 1 - Imagen y Texto Largo"),
    "carrusel": Path("Paquete 2 - Carrusel") / "Carrusel Publicación",
    "carrusel-historia": Path("Paquete 2 - Carrusel") / "Carrusel Historias",
    "texto-reflexivo": Path("Paquete 3 - Texto Reflexivo"),
}

CHECKLIST_ITEMS = [
    "[ ] Imagen revisada visualmente",
    "[ ] Copy revisado",
    "[ ] Primer comentario listo",
    "[ ] Publicado",
]


class PackagingError(RuntimeError):
    pass


def read_docx_paragraphs(path: Path) -> list:
    """Return one entry per logical line of the copy. Some callers save a
    whole multi-line caption as a single python-docx paragraph with literal
    "\\n" characters inside it rather than one paragraph object per line
    (seen in practice in an early carrusel-constelaciones caption) -- that
    collapses the real last line (the hashtags) into the middle of a giant
    paragraph instead of its own entry, which breaks the hashtag-detection
    heuristic below. Splitting every paragraph's text on "\\n" first makes
    both saving conventions produce the same flat line list."""
    if not path.exists():
        raise PackagingError(f"No existe --copy-docx: {path}")
    document = docx.Document(str(path))
    lines = []
    for p in document.paragraphs:
        for line in p.text.split("\n"):
            if line.strip():
                lines.append(line)
    return lines


def build_paquete(hook: str, copy_docx: Path, images: list, primer_comentario: Optional[str], out_dir: Path) -> Path:
    copy_paragraphs = read_docx_paragraphs(copy_docx)
    if not copy_paragraphs:
        raise PackagingError(f"{copy_docx} está vacío o no se pudo leer texto.")

    hashtags_line = None
    if copy_paragraphs[-1].strip().startswith("#"):
        hashtags_line = copy_paragraphs[-1].strip()

    doc = docx.Document()

    if images:
        doc.add_paragraph().add_run("IMÁGENES").bold = True
        for image_path in images:
            image_path = Path(image_path)
            if not image_path.exists():
                raise PackagingError(f"No existe --image: {image_path}")
            doc.add_picture(str(image_path), width=Inches(4))
            doc.add_paragraph().add_run(str(image_path)).italic = True

    doc.add_paragraph().add_run("COPY / CAPTION").bold = True
    for paragraph in copy_paragraphs:
        doc.add_paragraph(paragraph)

    if primer_comentario:
        doc.add_paragraph().add_run("PRIMER COMENTARIO (CTA)").bold = True
        doc.add_paragraph(primer_comentario)

    if hashtags_line:
        doc.add_paragraph().add_run("HASHTAGS").bold = True
        doc.add_paragraph(hashtags_line)

    doc.add_paragraph().add_run("CHECKLIST").bold = True
    for item in CHECKLIST_ITEMS:
        doc.add_paragraph(item)

    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"PAQUETE - {hook}.docx"
    doc.save(str(out_path))
    return out_path


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("hook", help="Hook o tema de la pieza (nombra el archivo final)")
    parser.add_argument("--copy-docx", required=True, help="Ruta al .docx con el copy/caption ya aprobado")
    parser.add_argument(
        "--image", action="append", default=[],
        help="Ruta a una imagen final, en orden de publicación. Repetible. Omitir para piezas sin imagen.",
    )
    parser.add_argument(
        "--primer-comentario", default=None,
        help="Texto corto (2-3 líneas) para fijar en el primer comentario, con "
        "su propio CTA al libro correcto. Opcional -- si se omite, esa "
        "sección del paquete no se genera (piezas sin CTA por diseño, ej. "
        "historias-constelaciones).",
    )
    parser.add_argument(
        "--tipo-pieza", choices=sorted(PIECE_TYPE_SUBFOLDERS),
        default=None,
        help="Tipo de pieza -- determina la subcarpeta fija dentro de "
        "<fecha> <micronicho>/. Obligatorio salvo que se pase --out-dir.",
    )
    parser.add_argument(
        "--micronicho", default=None,
        help="Slug corto del micronicho (ej. 'dolor-heredado'), reusado igual "
        "entre todas las piezas del mismo día/tema para que caigan en la "
        "misma carpeta. Obligatorio salvo que se pase --out-dir.",
    )
    parser.add_argument(
        "--fecha", default=None,
        help="Fecha YYYY-MM-DD para el nombre de la carpeta (default: hoy). "
        "Solo tiene efecto junto con --micronicho.",
    )
    parser.add_argument(
        "--out-dir", default=None,
        help="Override puntual de la carpeta de salida completa -- reemplaza "
        "el cálculo automático por --tipo-pieza/--micronicho/--fecha.",
    )
    args = parser.parse_args()

    if args.out_dir:
        out_dir = Path(args.out_dir)
    elif args.micronicho and args.tipo_pieza:
        fecha = args.fecha or datetime.date.today().isoformat()
        out_dir = BASE_PUBLICACIONES_DIR / f"{fecha} {args.micronicho}" / PIECE_TYPE_SUBFOLDERS[args.tipo_pieza]
    else:
        print(
            "Error: faltan --micronicho y/o --tipo-pieza (o --out-dir "
            "explícito) -- no se puede calcular la carpeta de salida.",
            file=sys.stderr,
        )
        return 1

    try:
        out_path = build_paquete(
            hook=args.hook,
            copy_docx=Path(args.copy_docx),
            images=args.image,
            primer_comentario=args.primer_comentario,
            out_dir=out_dir,
        )
    except PackagingError as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1

    print(f"Paquete guardado en: {out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
